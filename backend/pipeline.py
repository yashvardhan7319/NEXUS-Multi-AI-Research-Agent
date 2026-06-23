# pipeline.py — search -> scrape -> writer chain -> critic chain
# + standalone critic mode, compare mode, docx upload+download (pdf removed)

from tools import web_search, scrape_url
from agent import writer_chain, critic_chain, compare_chain
import re
import os
from datetime import datetime
import io
import asyncio
from docx import Document
from pypdf import PdfReader
from groq import RateLimitError


def run_research_pipeline(topic: str) -> dict:
    state = {}  # shared state dict, passed thru each step

    # ===== SEARCH AGENT =====
    print("\n" + "="*50)
    print("step 1 - search agent is working...")
    print("="*50)

    # direct tool call — keeps raw Title/URL/Snippet format intact, no LLM paraphrase
    state["search_results"] = web_search.invoke(topic)
    print(state["search_results"])

    # ===== READER AGENT =====
    print("\n" + "="*50)
    print("step 2 - reader agent is scraping multiple sources...")
    print("="*50)

    # grab ALL urls from search results, dedupe, cap at 5 (avoid rate-limit/cost blowup)
    urls = re.findall(r'URL:\s*(\S+)', state["search_results"])  # match tools.py output format exact
    seen = set()
    urls = [u for u in urls if not (u in seen or seen.add(u))]  # dedupe keep order
    urls = urls[:5]

    scraped_chunks = []
    if urls:
        for url in urls:
            print(f"  scraping: {url}")
            content = scrape_url.invoke(url)  # direct tool call, no agent wrapper
            scraped_chunks.append(f"SOURCE: {url}\n{content}")
    else:
        scraped_chunks.append("No URLs found.")

    state["scraped_content"] = "\n\n---\n\n".join(scraped_chunks)
    state["sources"] = urls  # keep raw url list for Sources section + report metadata
    print(state["scraped_content"])

    # ===== WRITER AGENT =====
    print("\n" + "="*50)
    print("step 3 - Writer is drafting the report...")
    print("="*50)

    # combine search + scraped data into single research blob for writer
    research_combined = (
        f"SEARCH RESULTS:\n{state['search_results']}\n\n"
        f"DETAILED SCRAPED CONTENT FROM MULTIPLE SOURCES:\n{state['scraped_content']}"
    )

    # writer_chain = prompt | llm | parser, builds full 12-section report
    try:
        state["report"] = writer_chain.invoke({
            "topic": topic,
            "research": research_combined
        })
    except RateLimitError as e:
        print("\n[ERROR] Groq daily token limit hit. Wait for reset or upgrade tier.")
        print(f"Details: {e}")
        raise SystemExit(1)

    print("\nFinal Report\n", state['report'])

    # ===== CRITIC AGENT =====
    print("\n" + "="*50)
    print("step 4 - critic is reviewing the report")
    print("="*50)

    # critic reviews final report, gives score + feedback
    try:
        state["feedback"] = critic_chain.invoke({
            "report": state['report']
        })
    except RateLimitError as e:
        print("\n[ERROR] Groq daily token limit hit during critic step.")
        print(f"Details: {e}")
        state["feedback"] = "[critic skipped — rate limit hit]"

    print("\ncritic report\n", state['feedback'])

    # final user-facing output — report already has its own ## 12. Sources section from writer prompt
    state["final_output"] = state["report"]

    return state


def review_user_report(report_text: str) -> str:
    """Feature: user uploads own report (docx/pdf), critic reviews it standalone.
    Bypasses search/scrape/writer steps entirely — straight to critic_chain."""
    # ===== CRITIC AGENT =====
    print("\n" + "="*50)
    print("critic is reviewing your uploaded report...")
    print("="*50)
    feedback = critic_chain.invoke({"report": report_text})
    print("\ncritic report\n", feedback)
    return feedback


def compare_reports(topic: str, report_a: str, report_b: str) -> str:
    """Feature: critic compares two reports (AI-generated vs user's own),
    scores both, picks a winner with reasoning."""
    # ===== CRITIC AGENT (compare mode) =====
    print("\n" + "="*50)
    print("critic is comparing both reports...")
    print("="*50)
    result = compare_chain.invoke({
        "topic": topic,
        "report_a": report_a,
        "report_b": report_b
    })
    print("\ncomparison result\n", result)
    return result


# ---------- file readers (docx/pdf only — upload still accepts both) ----------

def read_docx(filepath: str) -> str:
    # extract text from every non-empty paragraph, joined by newline
    doc = Document(filepath)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def read_pdf(filepath: str) -> str:
    # extract text page by page, join into single string
    reader = PdfReader(filepath)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def load_report_from_file(filepath: str) -> str:
    """Loads user-uploaded report. docx or pdf only — routes to correct reader by extension."""
    filepath = filepath.strip().strip('"').strip("'")  # strip wrapping quotes if user pastes quoted path
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".docx":
        return read_docx(filepath)
    elif ext == ".pdf":
        return read_pdf(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use .docx or .pdf only.")


# ---------- file writer (docx only — pdf download removed) ----------

def save_as_docx(content: str, filepath: str):
    # walk content line by line, map markdown-style headers (##, #) to docx heading styles
    doc = Document()
    for line in content.split("\n"):
        if line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        elif line.strip() == "":
            doc.add_paragraph("")  # preserve blank lines as spacing
        else:
            doc.add_paragraph(line)
    doc.save(filepath)


def save_report(content: str, topic: str) -> str:
    """Saves report as docx. Builds safe filename from topic + timestamp, writes file."""
    os.makedirs("reports", exist_ok=True)  # auto-create output dir if missing
    safe_topic = re.sub(r'[^\w\s-]', '', topic).strip().replace(' ', '_')[:50]  # strip unsafe filename chars
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filepath = f"reports/{safe_topic}_{timestamp}.docx"

    save_as_docx(content, filepath)

    print(f"\nReport saved to: {filepath}")
    return filepath


def ask_download() -> bool:
    # prompt user whether to download, returns True/False
    choice = input("Download this as docx? (y/n): ").strip().lower()
    return choice == "y"


def save_as_docx_to_bytes(content: str) -> io.BytesIO:
    doc = Document()
    for line in content.split("\n"):
        if line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        elif line.strip() == "":
            doc.add_paragraph("")  # preserve blank lines as spacing
        else:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


async def run_research_pipeline_stream(topic: str):
    state = {}

    # ===== SEARCH AGENT =====
    yield {"event": "step", "data": {"agent": "search", "status": "running", "content": "Search agent is querying Tavily..."}}
    try:
        # Run search in thread
        search_results = await asyncio.to_thread(web_search.invoke, topic)
        state["search_results"] = search_results
        yield {"event": "step", "data": {"agent": "search", "status": "done", "content": search_results}}
    except RateLimitError as e:
        yield {"event": "error", "data": {"message": "Groq daily token limit hit. Please try again later.", "retry_after": 60}}
        return
    except Exception as e:
        yield {"event": "error", "data": {"message": f"Search Agent failed: {str(e)}"}}
        return

    # ===== READER AGENT =====
    yield {"event": "step", "data": {"agent": "reader", "status": "running", "content": "Reader agent is extracting search links..."}}
    try:
        urls = re.findall(r'URL:\s*(\S+)', state["search_results"])
        seen = set()
        urls = [u for u in urls if not (u in seen or seen.add(u))]
        urls = urls[:5]

        scraped_chunks = []
        if urls:
            for url in urls:
                yield {"event": "step", "data": {"agent": "reader", "status": "running", "content": f"Scraping: {url}..."}}
                content = await asyncio.to_thread(scrape_url.invoke, url)
                scraped_chunks.append(f"SOURCE: {url}\n{content}")
        else:
            scraped_chunks.append("No URLs found.")

        state["scraped_content"] = "\n\n---\n\n".join(scraped_chunks)
        state["sources"] = urls
        yield {"event": "step", "data": {"agent": "reader", "status": "done", "content": f"Successfully scraped {len(urls)} sources:\n" + "\n".join(urls)}}
    except RateLimitError as e:
        yield {"event": "error", "data": {"message": "Groq daily token limit hit. Please try again later.", "retry_after": 60}}
        return
    except Exception as e:
        yield {"event": "error", "data": {"message": f"Reader Agent failed: {str(e)}"}}
        return

    # ===== WRITER AGENT =====
    yield {"event": "step", "data": {"agent": "writer", "status": "running", "content": "Writer agent is drafting the report..."}}
    research_combined = (
        f"SEARCH RESULTS:\n{state['search_results']}\n\n"
        f"DETAILED SCRAPED CONTENT FROM MULTIPLE SOURCES:\n{state['scraped_content']}"
    )
    try:
        report = await asyncio.to_thread(writer_chain.invoke, {
            "topic": topic,
            "research": research_combined
        })
        state["report"] = report
        yield {"event": "step", "data": {"agent": "writer", "status": "done", "content": report}}
    except RateLimitError as e:
        yield {"event": "error", "data": {"message": "Groq daily token limit hit. Please try again later.", "retry_after": 60}}
        return
    except Exception as e:
        yield {"event": "error", "data": {"message": f"Writer Agent failed: {str(e)}"}}
        return

    # ===== CRITIC AGENT =====
    yield {"event": "step", "data": {"agent": "critic", "status": "running", "content": "Critic agent is reviewing the report..."}}
    try:
        feedback = await asyncio.to_thread(critic_chain.invoke, {
            "report": state['report']
        })
        state["feedback"] = feedback
        yield {"event": "step", "data": {"agent": "critic", "status": "done", "content": feedback}}
    except RateLimitError as e:
        # Soft-fail or rate limit in critic
        yield {"event": "step", "data": {"agent": "critic", "status": "done", "content": "[critic skipped — rate limit hit]"}}
        state["feedback"] = "[critic skipped — rate limit hit]"
    except Exception as e:
        yield {"event": "step", "data": {"agent": "critic", "status": "done", "content": f"[critic skipped — error: {str(e)}]"}}
        state["feedback"] = f"[critic skipped — error: {str(e)}]"

    # ===== FINAL REPORT SSE EVENT =====
    yield {"event": "report", "data": {"report": state["report"], "feedback": state["feedback"]}}


if __name__ == "__main__":
    # ===== USER =====
    # CLI menu — 3 entry points into the system
    print("\nWhat do you want to do?")
    print("1. Generate new research report")
    print("2. Upload your own report  for critic review")
    print("3. Compare AI-generated report vs your own report")
    choice = input("\nEnter choice (1/2/3): ").strip()

    if choice == "1":
        # full pipeline: User -> Search Agent -> Reader Agent -> Writer Agent -> Critic Agent
        topic = input("\nEnter a research topic: ")
        result = run_research_pipeline(topic)
        print("\n\n" + "="*50)
        print("FINAL OUTPUT FOR USER")
        print("="*50)
        print(result["final_output"])

        if ask_download():
            save_report(result["final_output"], topic)

    elif choice == "2":
        # User -> Critic Agent (search/reader/writer skipped)
        filepath = input("\nEnter path to your report file : ").strip()
        report_text = load_report_from_file(filepath)
        feedback = review_user_report(report_text)

        if ask_download():
            save_report(feedback, "user_report_review")

    elif choice == "3":
        # User -> [Search Agent -> Reader Agent -> Writer Agent] -> Critic Agent (compare mode)
        topic = input("\nEnter the topic both reports cover: ")
        print("\n-- Generating AI report --")
        ai_result = run_research_pipeline(topic)

        filepath = input("\nEnter path to your own report file (.docx/.pdf): ").strip()
        user_report = load_report_from_file(filepath)

        comparison = compare_reports(topic, ai_result["final_output"], user_report)

        if ask_download():
            save_report(comparison, f"{topic}_comparison")

    else:
        print("Invalid choice.")